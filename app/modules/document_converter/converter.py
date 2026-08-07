"""
Document Converter - Core conversion functionality

Converts Markdown documents to DOCX and HTML formats with legal styling.
"""

import logging
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path

from app.core.utc import utc_now

logger = logging.getLogger(__name__)


@dataclass
class DocumentMetadata:
    """Metadata for converted documents."""
    title: str | None = None
    case_number: str | None = None
    court: str | None = None
    parties: str | None = None
    author: str | None = None
    created_at: datetime = None

    def __post_init__(self):
        if self.created_at is None:
            self.created_at = utc_now()


@dataclass
class DocumentStyle:
    """Styling configuration for document conversion."""
    name: str = "standard"
    font_family: str = "Times New Roman"
    font_size: int = 12
    line_spacing: float = 1.5
    margin_top: float = 1.0
    margin_bottom: float = 1.0
    margin_left: float = 1.0
    margin_right: float = 1.0


DOCUMENT_STYLES = {
    "standard": DocumentStyle(),
    "legal_brief": DocumentStyle(name="legal_brief", font_family="Times New Roman", font_size=12, line_spacing=2.0, margin_top=1.0, margin_bottom=1.0, margin_left=1.5, margin_right=1.0),
    "court_filing": DocumentStyle(name="court_filing", font_family="Courier New", font_size=12, line_spacing=2.0, margin_top=1.0, margin_bottom=1.0, margin_left=1.25, margin_right=1.25),
    "memo": DocumentStyle(name="memo", font_family="Arial", font_size=11, line_spacing=1.15, margin_top=1.0, margin_bottom=1.0, margin_left=1.0, margin_right=1.0),
}


class DocumentConverter:
    """Main document converter class."""

    def __init__(self):
        self.output_dir = Path("data/converted_documents")
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def convert_to_docx(
        self,
        markdown_text: str,
        metadata: DocumentMetadata | None = None,
        style: DocumentStyle | None = None
    ) -> Path:
        """Convert markdown text to DOCX format."""
        try:
            import docx
            from docx import Document
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
            from docx.shared import Inches, Pt
        except ImportError:
            raise ImportError("python-docx is required for DOCX conversion")

        if metadata is None:
            metadata = DocumentMetadata()
        if style is None:
            style = DocumentStyle.STYLES["standard"]

        # Create document
        doc = Document()

        # Set default font
        style_obj = doc.styles['Normal']
        font = style_obj.font
        font.name = style.font_family
        font.size = Pt(style.font_size)

        # Add title if provided
        if metadata.title:
            title_para = doc.add_paragraph(metadata.title)
            title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            title_run = title_para.runs[0]
            title_run.font.size = Pt(style.font_size + 4)
            title_run.font.bold = True

        # Add metadata header
        if any([metadata.case_number, metadata.court, metadata.parties, metadata.author]):
            meta_para = doc.add_paragraph()
            if metadata.case_number:
                meta_para.add_run(f"Case No: {metadata.case_number}\n")
            if metadata.court:
                meta_para.add_run(f"Court: {metadata.court}\n")
            if metadata.parties:
                meta_para.add_run(f"Parties: {metadata.parties}\n")
            if metadata.author:
                meta_para.add_run(f"Author: {metadata.author}\n")
            if metadata.created_at:
                meta_para.add_run(f"Date: {metadata.created_at.strftime('%B %d, %Y')}")

        # Add markdown content (basic conversion)
        lines = markdown_text.split('\n')
        for line in lines:
            if line.strip():
                if line.startswith('# '):
                    # H1 heading
                    para = doc.add_paragraph(line[2:])
                    para.style = 'Heading 1'
                elif line.startswith('## '):
                    # H2 heading
                    para = doc.add_paragraph(line[3:])
                    para.style = 'Heading 2'
                elif line.startswith('### '):
                    # H3 heading
                    para = doc.add_paragraph(line[4:])
                    para.style = 'Heading 3'
                else:
                    # Regular paragraph
                    doc.add_paragraph(line)

        # Save document
        filename = f"document_{utc_now().strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = self.output_dir / filename
        doc.save(filepath)

        logger.info(f"Converted markdown to DOCX: {filepath}")
        return filepath

    def convert_to_html(
        self,
        markdown_text: str,
        metadata: DocumentMetadata | None = None,
        style: DocumentStyle | None = None
    ) -> Path:
        """Convert markdown text to interactive HTML format."""
        try:
            import markdown
        except ImportError:
            raise ImportError("markdown is required for HTML conversion")

        if metadata is None:
            metadata = DocumentMetadata()
        if style is None:
            style = DocumentStyle.STYLES["standard"]

        # Convert markdown to HTML
        html_content = markdown.markdown(markdown_text, extensions=['tables', 'fenced_code'])

        # Create HTML document
        html_template = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{metadata.title or 'Document'}</title>
    <style>
        body {{
            font-family: {style.font_family};
            font-size: {style.font_size}pt;
            line-height: {style.line_spacing};
            margin: {style.margin_top}in {style.margin_right}in {style.margin_bottom}in {style.margin_left}in;
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
        }}
        h1 {{ font-size: {style.font_size + 6}pt; margin-top: 24px; margin-bottom: 12px; }}
        h2 {{ font-size: {style.font_size + 4}pt; margin-top: 20px; margin-bottom: 10px; }}
        h3 {{ font-size: {style.font_size + 2}pt; margin-top: 16px; margin-bottom: 8px; }}
        .metadata {{
            background: #f5f5f5;
            padding: 10px;
            border-left: 4px solid #ccc;
            margin-bottom: 20px;
        }}
        .metadata p {{ margin: 5px 0; }}
        table {{ border-collapse: collapse; width: 100%; margin: 15px 0; }}
        th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
        th {{ background-color: #f2f2f2; }}
        code {{ background: #f4f4f4; padding: 2px 4px; border-radius: 3px; }}
        pre {{ background: #f4f4f4; padding: 10px; border-radius: 5px; overflow-x: auto; }}
    </style>
</head>
<body>
"""

        # Add metadata header
        if any([metadata.title, metadata.case_number, metadata.court, metadata.parties, metadata.author]):
            html_template += '<div class="metadata">'
            if metadata.title:
                html_template += f'<h1>{metadata.title}</h1>'
            if metadata.case_number:
                html_template += f'<p><strong>Case No:</strong> {metadata.case_number}</p>'
            if metadata.court:
                html_template += f'<p><strong>Court:</strong> {metadata.court}</p>'
            if metadata.parties:
                html_template += f'<p><strong>Parties:</strong> {metadata.parties}</p>'
            if metadata.author:
                html_template += f'<p><strong>Author:</strong> {metadata.author}</p>'
            if metadata.created_at:
                html_template += f'<p><strong>Date:</strong> {metadata.created_at.strftime("%B %d, %Y")}</p>'
            html_template += '</div>'

        # Add content
        html_template += f'<div class="content">{html_content}</div>'
        html_template += '</body></html>'

        # Save HTML file
        filename = f"document_{utc_now().strftime('%Y%m%d_%H%M%S')}.html"
        filepath = self.output_dir / filename

        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(html_template)

        logger.info(f"Converted markdown to HTML: {filepath}")
        return filepath


# Convenience functions
def markdown_to_docx(
    markdown_text: str,
    title: str | None = None,
    case_number: str | None = None,
    court: str | None = None,
    parties: str | None = None,
    author: str | None = None,
    style: str = "standard"
) -> Path:
    """Convert markdown to DOCX with metadata."""
    metadata = DocumentMetadata(
        title=title,
        case_number=case_number,
        court=court,
        parties=parties,
        author=author
    )

    doc_style = DocumentStyle.STYLES.get(style, DocumentStyle.STYLES["standard"])

    converter = DocumentConverter()
    return converter.convert_to_docx(markdown_text, metadata, doc_style)


def markdown_to_html(
    markdown_text: str,
    title: str | None = None,
    case_number: str | None = None,
    court: str | None = None,
    parties: str | None = None,
    author: str | None = None,
    style: str = "standard"
) -> Path:
    """Convert markdown to HTML with metadata."""
    metadata = DocumentMetadata(
        title=title,
        case_number=case_number,
        court=court,
        parties=parties,
        author=author
    )

    doc_style = DocumentStyle.STYLES.get(style, DocumentStyle.STYLES["standard"])

    converter = DocumentConverter()
    return converter.convert_to_html(markdown_text, metadata, doc_style)
